import re
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_word_runs(paragraph, text: str):
    """Parse **bold** and normal text, add as runs to a Word paragraph."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_section_content(doc: Document, content: str):
    """Add section content, handling bullet lists and markdown."""
    if not content:
        return
    lines = content.split('\n')
    i = 0
    buffer = []

    while i < len(lines):
        line = lines[i]

        # Bullet list item
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            if buffer:
                _flush_buffer(doc, buffer)
                buffer = []
            bullet_text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            markdown_to_word_runs(p, bullet_text)
        elif line.strip() == '':
            if buffer:
                _flush_buffer(doc, buffer)
                buffer = []
        else:
            # Strip common markdown: headings, code blocks, horizontal rules
            clean = re.sub(r'^#{1,6}\s*', '', line)
            clean = re.sub(r'`{1,3}[^`]*`{1,3}', lambda m: m.group().strip('`'), clean)
            clean = re.sub(r'^\*{3,}$|^-{3,}$|^_{3,}$', '', clean)
            if clean.strip():
                buffer.append(clean)
        i += 1

    if buffer:
        _flush_buffer(doc, buffer)


def _flush_buffer(doc: Document, lines: list[str]):
    text = ' '.join(lines)
    p = doc.add_paragraph()
    markdown_to_word_runs(p, text)


def export_project_to_docx(project, sections_data: list[dict]) -> BytesIO:
    doc = Document()

    heading_style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}

    for sec_info in sections_data:
        eff = sec_info["effective_section"]
        content = sec_info.get("content")
        level = eff.get("level", 1)
        title = eff.get("title", "")

        style = heading_style_map.get(level, 'Heading 1')
        doc.add_heading(title, level=level)

        if content and content.strip():
            add_section_content(doc, content)
        else:
            p = doc.add_paragraph("[待填写]")
            p.runs[0].italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
